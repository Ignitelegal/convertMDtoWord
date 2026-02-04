#!/usr/bin/env python3
"""
Markdown to Word Converter with Template Styling
================================================
Converts markdown (.md) files to Word (.docx) documents while applying
styles from a template file for consistent organizational formatting.

Author: Generated for Australian Unity Legal & Governance
Date: 2026-02-04
"""

import argparse
import sys
import re
from pathlib import Path
from typing import Optional, List, Tuple
import mistune
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWordConverter:
    """
    Converts markdown files to Word documents using template styling.
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the converter with an optional template.

        Args:
            template_path: Path to Word template file (.docx)
        """
        self.template_path = template_path
        self.doc = None
        self.current_list_level = 0
        self.in_code_block = False

        # Style mapping: markdown element -> Word style name
        self.style_map = {
            'h1': 'Heading 1',
            'h2': 'Heading 2',
            'h3': 'Heading 3',
            'h4': 'Heading 4',
            'h5': 'Heading 5',
            'h6': 'Heading 6',
            'body': 'Body Text',
            'paragraph': 'Normal',
            'bullet': 'List Bullet',
            'numbered': 'List Number',
            'code_block': 'Code',
            'quote': 'Quote',
            'intense_quote': 'Intense Quote'
        }

    def load_template(self) -> Document:
        """
        Load Word template or create new document.

        Returns:
            Document object
        """
        try:
            if self.template_path and Path(self.template_path).exists():
                print(f"📄 Loading template: {self.template_path}")
                return Document(self.template_path)
            else:
                if self.template_path:
                    print(f"⚠️  Template not found: {self.template_path}")
                    print("📄 Creating new document with default styles")
                return Document()
        except Exception as e:
            print(f"❌ Error loading template: {e}")
            print("📄 Creating new document with default styles")
            return Document()

    def get_style_or_fallback(self, preferred_style: str, fallback: str = 'Normal') -> str:
        """
        Get style name from document or fallback to default.

        Args:
            preferred_style: Preferred style name
            fallback: Fallback style name

        Returns:
            Available style name
        """
        available_styles = [s.name for s in self.doc.styles]

        if preferred_style in available_styles:
            return preferred_style

        # Try variations
        variations = [
            preferred_style.replace(' ', ''),
            preferred_style.replace(' ', '_'),
            preferred_style.lower(),
            preferred_style.upper()
        ]

        for variation in variations:
            if variation in available_styles:
                return variation

        # Fallback
        if fallback in available_styles:
            return fallback

        return 'Normal'

    def read_markdown_file(self, md_path: str) -> str:
        """
        Read markdown file content.

        Args:
            md_path: Path to markdown file

        Returns:
            File content as string

        Raises:
            FileNotFoundError: If file doesn't exist
            UnicodeDecodeError: If file encoding is invalid
        """
        md_file = Path(md_path)

        if not md_file.exists():
            raise FileNotFoundError(f"Markdown file not found: {md_path}")

        if not md_file.suffix.lower() in ['.md', '.markdown']:
            print(f"⚠️  Warning: File doesn't have .md extension: {md_path}")

        try:
            content = md_file.read_text(encoding='utf-8')
            print(f"✅ Read {len(content)} characters from {md_file.name}")
            return content
        except UnicodeDecodeError:
            # Try alternative encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    content = md_file.read_text(encoding=encoding)
                    print(f"✅ Read file using {encoding} encoding")
                    return content
                except UnicodeDecodeError:
                    continue
            raise UnicodeDecodeError(
                'utf-8', b'', 0, 1, 
                f"Unable to decode file with any standard encoding"
            )

    def parse_inline_formatting(self, text: str, paragraph):
        """
        Parse and apply inline formatting (bold, italic, code, links).

        Args:
            text: Text with markdown inline formatting
            paragraph: Document paragraph object to add runs to
        """
        # Pattern for inline elements: **bold**, *italic*, `code`, [link](url)
        # This regex handles nested and mixed formatting
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))'

        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            run = paragraph.add_run()

            # Bold: **text**
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True

            # Italic: *text* (but not if it's part of **)
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run.text = part[1:-1]
                run.italic = True

            # Inline code: `text`
            elif part.startswith('`') and part.endswith('`'):
                run.text = part[1:-1]
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

            # Link: [text](url)
            elif part.startswith('[') and '](' in part:
                match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if match:
                    link_text, url = match.groups()
                    run.text = link_text
                    # Add hyperlink styling
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True

            # Plain text
            else:
                run.text = part

    def add_heading(self, text: str, level: int):
        """
        Add heading to document.

        Args:
            text: Heading text
            level: Heading level (1-6)
        """
        style_key = f'h{level}'
        style_name = self.get_style_or_fallback(
            self.style_map.get(style_key, 'Normal'),
            'Normal'
        )

        # Remove markdown heading markers
        text = re.sub(r'^#{1,6}\s+', '', text).strip()

        paragraph = self.doc.add_paragraph(style=style_name)
        self.parse_inline_formatting(text, paragraph)

    def add_paragraph(self, text: str, style: str = 'Normal'):
        """
        Add paragraph to document.

        Args:
            text: Paragraph text
            style: Word style name
        """
        if not text.strip():
            self.doc.add_paragraph()
            return

        style_name = self.get_style_or_fallback(style, 'Normal')
        paragraph = self.doc.add_paragraph(style=style_name)
        self.parse_inline_formatting(text, paragraph)

    def add_list_item(self, text: str, ordered: bool = False, level: int = 0):
        """
        Add list item to document.

        Args:
            text: List item text
            ordered: True for numbered list, False for bullet list
            level: Nesting level (0-based)
        """
        style_key = 'numbered' if ordered else 'bullet'
        style_name = self.get_style_or_fallback(
            self.style_map.get(style_key, 'Normal'),
            'Normal'
        )

        # Remove markdown list markers
        text = re.sub(r'^[\*\-\+]\s+', '', text)  # Unordered
        text = re.sub(r'^\d+\.\s+', '', text)     # Ordered

        paragraph = self.doc.add_paragraph(style=style_name)

        # Apply indentation for nested lists
        if level > 0:
            paragraph.paragraph_format.left_indent = Inches(0.5 * level)

        self.parse_inline_formatting(text, paragraph)

    def add_code_block(self, code: str, language: str = ''):
        """
        Add code block to document.

        Args:
            code: Code content
            language: Programming language (for reference)
        """
        style_name = self.get_style_or_fallback(
            self.style_map.get('code_block', 'Normal'),
            'Normal'
        )

        # Split code into lines and add each as separate paragraph
        lines = code.split('
')

        for line in lines:
            paragraph = self.doc.add_paragraph(line, style=style_name)

            # Apply code formatting if style doesn't exist
            if style_name == 'Normal':
                for run in paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

                # Add light gray background shading
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F0F0F0')
                paragraph._element.get_or_add_pPr().append(shading_elm)

    def add_blockquote(self, text: str):
        """
        Add blockquote to document.

        Args:
            text: Quote text
        """
        style_name = self.get_style_or_fallback(
            self.style_map.get('quote', 'Normal'),
            'Normal'
        )

        # Remove markdown quote marker
        text = re.sub(r'^>\s+', '', text)

        paragraph = self.doc.add_paragraph(style=style_name)

        # Apply quote formatting if style doesn't exist
        if style_name == 'Normal':
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.right_indent = Inches(0.5)

        self.parse_inline_formatting(text, paragraph)

    def add_horizontal_rule(self):
        """Add horizontal rule (line) to document."""
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)

        # Add bottom border to create horizontal line
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def process_markdown_lines(self, md_content: str):
        """
        Process markdown content line by line.

        Args:
            md_content: Markdown file content
        """
        lines = md_content.split('
')
        i = 0
        in_code_block = False
        code_buffer = []
        code_language = ''
        in_list = False

        while i < len(lines):
            line = lines[i]

            # Code block detection
            if line.strip().startswith('```'):
                if not in_code_block:
                    # Start code block
                    in_code_block = True
                    code_language = line.strip()[3:].strip()
                    code_buffer = []
                else:
                    # End code block
                    in_code_block = False
                    self.add_code_block('
'.join(code_buffer), code_language)
                    code_buffer = []
                    code_language = ''
                i += 1
                continue

            # Inside code block
            if in_code_block:
                code_buffer.append(line)
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^[-*_]{3,}$', line.strip()):
                self.add_horizontal_rule()
                i += 1
                continue

            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self.add_heading(text, level)
                in_list = False
                i += 1
                continue

            # Blockquote
            if line.strip().startswith('>'):
                self.add_blockquote(line)
                in_list = False
                i += 1
                continue

            # Ordered list
            ordered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
            if ordered_match:
                indent = len(ordered_match.group(1))
                text = ordered_match.group(2)
                level = indent // 2
                self.add_list_item(text, ordered=True, level=level)
                in_list = True
                i += 1
                continue

            # Unordered list
            unordered_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
            if unordered_match:
                indent = len(unordered_match.group(1))
                text = unordered_match.group(2)
                level = indent // 2
                self.add_list_item(text, ordered=False, level=level)
                in_list = True
                i += 1
                continue

            # Empty line
            if not line.strip():
                if in_list:
                    in_list = False
                self.doc.add_paragraph()
                i += 1
                continue

            # Regular paragraph
            self.add_paragraph(line)
            in_list = False
            i += 1

    def convert(self, md_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert markdown file to Word document.

        Args:
            md_path: Path to markdown file
            output_path: Path for output Word file (optional)

        Returns:
            Path to generated Word document

        Raises:
            FileNotFoundError: If markdown file not found
            PermissionError: If unable to write output file
        """
        print("
" + "="*60)
        print("MARKDOWN TO WORD CONVERTER")
        print("="*60 + "
")

        # Read markdown file
        print("📖 Reading markdown file...")
        md_content = self.read_markdown_file(md_path)

        # Load template
        print("
📋 Loading template...")
        self.doc = self.load_template()

        # List available styles for debugging
        available_styles = [s.name for s in self.doc.styles]
        print(f"✅ Template loaded with {len(available_styles)} styles")

        # Process markdown
        print("
🔄 Converting markdown to Word...")
        self.process_markdown_lines(md_content)

        # Determine output path
        if not output_path:
            md_file = Path(md_path)
            output_path = md_file.parent / f"{md_file.stem}_converted.docx"

        output_file = Path(output_path)

        # Save document
        print(f"
💾 Saving document to: {output_file}")
        try:
            self.doc.save(str(output_file))
            print(f"✅ Conversion successful!")
            print(f"📄 Output: {output_file.absolute()}")
            return str(output_file)
        except PermissionError:
            raise PermissionError(
                f"Unable to write to {output_file}. "
                "Check if file is open in another program."
            )
        except Exception as e:
            raise Exception(f"Error saving document: {e}")


def main():
    """Main execution function for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Convert markdown files to Word documents with template styling',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic conversion
  python md_to_word.py document.md

  # With template
  python md_to_word.py document.md -t template.docx

  # Specify output location
  python md_to_word.py document.md -o output/report.docx

  # With template and custom output
  python md_to_word.py document.md -t template.docx -o final_report.docx
        """
    )

    parser.add_argument(
        'markdown_file',
        help='Path to the markdown (.md) file to convert'
    )

    parser.add_argument(
        '-t', '--template',
        help='Path to Word template (.docx) file for styling',
        default=None
    )

    parser.add_argument(
        '-o', '--output',
        help='Path for output Word document (default: same location as input)',
        default=None
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output'
    )

    args = parser.parse_args()

    try:
        # Create converter
        converter = MarkdownToWordConverter(template_path=args.template)

        # Convert file
        output_path = converter.convert(args.markdown_file, args.output)

        print("
" + "="*60)
        print("CONVERSION COMPLETE")
        print("="*60)
        print(f"
✅ Successfully converted:")
        print(f"   Input:  {args.markdown_file}")
        print(f"   Output: {output_path}
")

        return 0

    except FileNotFoundError as e:
        print(f"
❌ ERROR: {e}", file=sys.stderr)
        return 1

    except PermissionError as e:
        print(f"
❌ ERROR: {e}", file=sys.stderr)
        return 1

    except Exception as e:
        print(f"
❌ UNEXPECTED ERROR: {e}", file=sys.stderr)
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
